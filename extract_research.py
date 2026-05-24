import docx
import sys

# Set UTF-8 encoding for output
sys.stdout.reconfigure(encoding='utf-8')

# Get filename from command line argument
filename = sys.argv[1] if len(sys.argv) > 1 else 'FINAL RP-Draft.docx'

try:
    doc = docx.Document(filename)
    
    # Extract all paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            print(paragraph.text)
    
    # Extract tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    print(cell.text)
except Exception as e:
    print(f"Error: {e}")

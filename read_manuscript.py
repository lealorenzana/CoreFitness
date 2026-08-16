"""
Read the COREFITNESS.docx manuscript and extract its contents.
Then generate a Q&A based on the manuscript content.
"""

from docx import Document
import os

def read_docx(filepath):
    """Read all text content from a .docx file."""
    doc = Document(filepath)
    content = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text.strip())
    
    # Also read tables if any
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                content.append(row_text)
    
    return '\n'.join(content)

def main():
    filepath = r"C:\Users\ASUS\Desktop\CoreFitness\docs\COREFITNESS.docx"
    
    if not os.path.exists(filepath):
        print(f"ERROR: File not found at {filepath}")
        return
    
    print("=" * 80)
    print("READING MANUSCRIPT: COREFITNESS.docx")
    print("=" * 80)
    
    content = read_docx(filepath)
    
    # Save the extracted content to a text file for reference
    output_path = r"C:\Users\ASUS\Desktop\CoreFitness\docs\manuscript_content.txt"
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"\nManuscript content extracted and saved to: {output_path}")
    print(f"Total characters: {len(content)}")
    print(f"Total lines: {len(content.splitlines())}")
    print("\n" + "=" * 80)
    print("MANUSCRIPT CONTENT PREVIEW (first 5000 chars):")
    print("=" * 80)
    print(content[:5000])
    print("\n..." if len(content) > 5000 else "")
    print("=" * 80)

if __name__ == "__main__":
    main()

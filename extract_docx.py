import docx
import sys

# Get filename from command line argument
filename = sys.argv[1] if len(sys.argv) > 1 else 'FINAL RP-Draft.docx'
doc = docx.Document(filename)

# Extract all paragraphs
for paragraph in doc.paragraphs:
    print(paragraph.text.encode('utf-8', errors='ignore').decode('utf-8'))

# Extract tables if any
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                print(cell.text.encode('utf-8', errors='ignore').decode('utf-8'))
